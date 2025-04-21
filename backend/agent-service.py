import os
import numpy as np
from fastapi import FastAPI, File, UploadFile, Form
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import List
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from dotenv import load_dotenv
from openai import OpenAI
from langchain.text_splitter import CharacterTextSplitter
from langchain.schema import Document
from langchain_community.embeddings import OpenAIEmbeddings
from langchain_community.vectorstores.supabase import SupabaseVectorStore
from supabase import create_client, Client

load_dotenv()

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Adjust for production
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- Config ---
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_KEY = os.getenv("SUPABASE_KEY")
supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)

# --- Custom Vector Store for UUID-based documents ---
class SupabaseUUIDVectorStore(SupabaseVectorStore):
    def __init__(self, client, embedding, table_name):
        self.client = client
        self.embedding = embedding
        self._embedding = embedding
        self.table_name = table_name

    @classmethod
    def from_texts(cls, texts, embedding, client, table_name, metadata=None):
        vectors = embedding.embed_documents(texts)
        rows = []

        for i, text in enumerate(texts):
            row = {
                "content": text,
                "embedding": vectors[i],
            }
            if metadata:
                row["metadata"] = metadata[i]
            rows.append(row)

        insert_resp = client.table(table_name).insert(rows).execute()
        if insert_resp.data is None:
            raise Exception(f"Failed to insert documents")

        return cls(client=client, embedding=embedding, table_name=table_name)

    def similarity_search_by_vector_with_relevance_scores(self, embedding, k=4, filter=None, **kwargs):
        if isinstance(embedding, np.ndarray):
            embedding = embedding.tolist()

        try:
            response = self.client.rpc(
                "match_documents",
                {
                    "query_embedding": embedding,
                    "match_count": k,
                    "filter": filter or {},
                },
            ).execute()
        except Exception as e:
            raise ValueError(f"Error querying Supabase: {e}")

        if not response.data:
            return []

        docs = []
        scores = []
        for match in response.data:
            metadata = match.get("metadata", {})
            text = match.get("content", "")
            score = match.get("similarity", 0)
            docs.append(Document(page_content=text, metadata=metadata))
            scores.append(score)

        return list(zip(docs, scores))

# --- Helpers ---
def extract_text(file: UploadFile):
    if file.filename.endswith(".pdf"):
        reader = PdfReader(file.file)
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    elif file.filename.endswith(".docx"):
        doc = DocxDocument(file.file)
        return "\n".join(p.text for p in doc.paragraphs)
    return ""

def process_text(text: str):
    splitter = CharacterTextSplitter(chunk_size=10000, chunk_overlap=50)
    return splitter.split_text(text)

def store_embeddings(chunks):
    embeddings = OpenAIEmbeddings()
    return SupabaseUUIDVectorStore.from_texts(chunks, embeddings, client=supabase, table_name="documents")

def chat_with_doc(db, query: str):
    results = db.similarity_search(query, k=3)
    context = "\n".join([r.page_content for r in results])
    prompt = f"Answer the question based on the following context:\n\n{context}\n\nQuestion: {query}"
    response = client.chat.completions.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": prompt}
        ]
    )
    return response.choices[0].message.content

# --- API Models ---
class ChatInput(BaseModel):
    query: str

class Message(BaseModel):
    role: str
    content: str
    
class ChatRequest(BaseModel):
    model: str
    messages: list[Message]


# --- In-Memory Store ---
memory_db = {}

# --- Routes ---
@app.post("/v1/upload")
async def upload(file: UploadFile = File(...)):
    text = extract_text(file)
    chunks = process_text(text)
    db = store_embeddings(chunks)
    memory_db["db"] = db
    return {"message": "File processed and embedded successfully."}

@app.post("/v1/chat")
def chat(input: ChatInput):
    db = memory_db.get("db")
    if not db:
        return {"error": "No document has been uploaded yet."}
    answer = chat_with_doc(db, input.query)
    return {"response": answer}

@app.post("/v1/chat/completions")
async def chat_completion(req: ChatRequest):
    db = memory_db.get("db")
    if not db:
        return {"error": "No document has been uploaded yet."}
    user_query = req.messages[-1].content  # Last message from user
    answer = chat_with_doc(db, user_query)
    return {
        "id": "chatcmpl-custom",
        "object": "chat.completion",
        "choices": [{
            "message": {
                "role": "assistant",
                "content": answer
            },
            "finish_reason": "stop",
            "index": 0
        }],
        "usage": {}
    }

